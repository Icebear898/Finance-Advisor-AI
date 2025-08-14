import PyPDF2
import openpyxl
from docx import Document
import os
import logging
from typing import Optional
import pandas as pd

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Utility class for processing different document formats"""
    
    def __init__(self):
        self.supported_formats = {
            'pdf': self._extract_pdf_text,
            'docx': self._extract_docx_text,
            'xlsx': self._extract_xlsx_text,
            'xls': self._extract_xlsx_text,
            'txt': self._extract_txt_text
        }

    async def extract_text(self, file_path: str, file_type: str) -> Optional[str]:
        """Extract text content from a file based on its type"""
        try:
            if not os.path.exists(file_path):
                logger.error(f"File not found: {file_path}")
                return None
            
            file_type = file_type.lower()
            
            if file_type not in self.supported_formats:
                logger.error(f"Unsupported file type: {file_type}")
                return None
            
            # Extract text using appropriate method
            extractor = self.supported_formats[file_type]
            content = await extractor(file_path)
            
            if content:
                logger.info(f"Successfully extracted {len(content)} characters from {file_path}")
            
            return content
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return None

    async def _extract_pdf_text(self, file_path: str) -> Optional[str]:
        """Extract text from PDF file"""
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    try:
                        page = pdf_reader.pages[page_num]
                        text = page.extract_text()
                        if text.strip():
                            text_content.append(f"Page {page_num + 1}:\n{text}")
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                        continue
            
            return "\n\n".join(text_content) if text_content else None
            
        except Exception as e:
            logger.error(f"Error processing PDF file {file_path}: {str(e)}")
            return None

    async def _extract_docx_text(self, file_path: str) -> Optional[str]:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                if table_text:
                    text_content.append("Table:\n" + "\n".join(table_text))
            
            return "\n\n".join(text_content) if text_content else None
            
        except Exception as e:
            logger.error(f"Error processing DOCX file {file_path}: {str(e)}")
            return None

    async def _extract_xlsx_text(self, file_path: str) -> Optional[str]:
        """Extract text from Excel file"""
        try:
            text_content = []
            
            # Read Excel file
            excel_file = pd.ExcelFile(file_path)
            
            for sheet_name in excel_file.sheet_names:
                try:
                    # Read sheet as DataFrame
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    if not df.empty:
                        sheet_content = [f"Sheet: {sheet_name}"]
                        
                        # Add column headers
                        headers = list(df.columns)
                        sheet_content.append("Headers: " + " | ".join(str(h) for h in headers))
                        
                        # Add data rows (limit to first 100 rows to avoid huge output)
                        for index, row in df.head(100).iterrows():
                            row_text = " | ".join(str(cell) for cell in row.values if pd.notna(cell))
                            if row_text.strip():
                                sheet_content.append(f"Row {index + 1}: {row_text}")
                        
                        text_content.append("\n".join(sheet_content))
                        
                except Exception as e:
                    logger.warning(f"Error processing sheet {sheet_name}: {str(e)}")
                    continue
            
            return "\n\n".join(text_content) if text_content else None
            
        except Exception as e:
            logger.error(f"Error processing Excel file {file_path}: {str(e)}")
            return None

    async def _extract_txt_text(self, file_path: str) -> Optional[str]:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                return content if content.strip() else None
                
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    content = file.read()
                    return content if content.strip() else None
            except Exception as e:
                logger.error(f"Error reading text file {file_path} with latin-1 encoding: {str(e)}")
                return None
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {str(e)}")
            return None

    def get_file_info(self, file_path: str) -> dict:
        """Get basic information about a file"""
        try:
            if not os.path.exists(file_path):
                return {"error": "File not found"}
            
            file_stats = os.stat(file_path)
            file_name = os.path.basename(file_path)
            file_extension = os.path.splitext(file_name)[1].lower()
            
            return {
                "filename": file_name,
                "file_size": file_stats.st_size,
                "file_extension": file_extension,
                "created_time": file_stats.st_ctime,
                "modified_time": file_stats.st_mtime,
                "is_readable": os.access(file_path, os.R_OK)
            }
            
        except Exception as e:
            logger.error(f"Error getting file info for {file_path}: {str(e)}")
            return {"error": str(e)}

    def validate_file_format(self, file_path: str) -> bool:
        """Validate if file format is supported"""
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            supported_extensions = ['.pdf', '.docx', '.xlsx', '.xls', '.txt']
            
            return file_extension in supported_extensions
            
        except Exception as e:
            logger.error(f"Error validating file format for {file_path}: {str(e)}")
            return False

    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = ' '.join(text.split())
        
        # Remove special characters that might cause issues
        import re
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\{\}]', '', text)
        
        return text.strip()
